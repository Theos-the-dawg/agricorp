from django.shortcuts import render

# Create your views here.
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .forms import (AddNewUserForm, LoginForm, MyUserCreationForm,ChangeUserDetailForm)
from .models import UserDetails,MyUser
from .utils import authenticate_custom_user
from django.http import HttpResponse
from .models import MyUser, UserDetails
from docx import Document
from django.views.decorators.csrf import csrf_exempt

@csrf_exempt
def home_view(request):
    # Example view for home page after successful login
    # user_id = request.session.get('user_id')
    # if user_id is None:
    #     return redirect('login')  # Redirect to login if user is not authenticated
    # user = MyUser.objects.get(pk=user_id)
    return render(request, 'home.html')#, {'user': user})

def login_view(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate_custom_user(request, username, password) 
            if user is not None:
                request.session['user_id'] = user.id  # Example: Storing user ID in session
                return redirect('home')
            else:
                messages.error(request, 'Invalid username or password')
        else:
            messages.error(request, 'Please correct the error below.')
    else:
        form = LoginForm()
    return render(request, 'login.html', {'form': form})

@csrf_exempt
def add_user_view(request):
 
    if request.method == 'POST':
        user_form = MyUserCreationForm(request.POST)
        detail_form = AddNewUserForm(request.POST)
        
        if user_form.is_valid() and detail_form.is_valid():
            
            #creates instance but doesn't save to database model
            user = user_form.save(commit=False)
            #check that values dont match existing error # debug in user form shows error handling
            user.set_password(user_form.cleaned_data['password1']) # password is set as password1 and password2 s
            
            user.save()
           
            user_detail = detail_form.save(commit=False)
            user_detail.user = user


            user_detail.facilitator = bool(int(detail_form.cleaned_data['facilitator']))
            user_detail.save()
            
            messages.success(request, 'User and user details added successfully!')
            return redirect('home')
        else:
            
            messages.error(request, 'Please correct the errors below.')
    else:
        user_form = MyUserCreationForm()
        detail_form = AddNewUserForm()
    
    context = {
        'user_form': user_form,
        'detail_form': detail_form,
    }
    
    return render(request, 'add_user.html', context)

def manage_user_view(request):

    users = UserDetails.objects.all()

    context={
        'users':users,
    }
    return render(request, 'manage_user.html',context)

@csrf_exempt
def edit_user_view(request, id):
    

    if request.method == 'POST':
     user = get_object_or_404(MyUser, id=id)
     user_detail = UserDetails.objects.get(user=user)

   
    change_form = ChangeUserDetailForm(request.POST, instance=user_detail)
    if change_form.is_valid():
        change_form.save()
        return redirect('manage_user')
       
    else:
        
        change_form = AddNewUserForm()
    context = {
        
        'change_form': change_form,
    }
    return render(request, 'edit_user.html', context)

def delete_user_view(request, id):
    user = get_object_or_404(MyUser, id=id)

    if request.method == 'POST':
        if 'confirm' in request.POST:
            user.delete()
            messages.success(request, 'User deleted successfully.')
            return redirect('manage_user')
        else:
            return redirect('manage_user')

    context = {
        'user': user,
    }

    return render(request, 'delete_user.html', context)

def export_to_word(request):
    users = UserDetails.objects.all()
    document = Document()
    document.add_heading('User Details', level=1)

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Date of Birth'
    hdr_cells[2].text = 'Province'
    hdr_cells[3].text = 'Gender'
    hdr_cells[4].text = 'Facilitator'

    for user in users:
        row_cells = table.add_row().cells
        row_cells[0].text = user.user.username
        row_cells[1].text = str(user.dateofbirth)
        row_cells[2].text = user.province
        row_cells[3].text = user.gender  
        row_cells[4].text = 'Yes' if user.facilitator else 'No'

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=user_details.docx'
    document.save(response)
    return response
